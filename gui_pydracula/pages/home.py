from __future__ import annotations

from pathlib import Path
from datetime import datetime
import json
from PySide6 import QtCore, QtWidgets, QtGui, QtMultimedia, QtMultimediaWidgets
from pipeline.helpers import preview_voice, sanitize_name
from pipeline.config import Config
from pipeline import generator
import subprocess


class ScriptEdit(QtWidgets.QPlainTextEdit):
    fileDropped = QtCore.Signal(str)

    def __init__(self, parent: QtWidgets.QWidget | None = None) -> None:
        super().__init__(parent)
        self.setAcceptDrops(True)

    def dragEnterEvent(self, event: QtGui.QDragEnterEvent) -> None:  # type: ignore[override]
        if event.mimeData().hasUrls():
            self.setStyleSheet("background:rgba(59,130,246,50);")
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)

    def dragLeaveEvent(self, event: QtGui.QDragLeaveEvent) -> None:  # type: ignore[override]
        self.setStyleSheet("")
        super().dragLeaveEvent(event)

    def dropEvent(self, event: QtGui.QDropEvent) -> None:  # type: ignore[override]
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path.lower().endswith(('.txt', '.docx')):
                self.fileDropped.emit(path)
                break
        else:
            super().dropEvent(event)
        self.setStyleSheet("")
        event.acceptProposedAction()


class HomePage(QtWidgets.QWidget):
    generateRequested = QtCore.Signal(dict)
    configChanged = QtCore.Signal(dict)

    def __init__(self, config: Config, parent: QtWidgets.QWidget | None = None) -> None:
        super().__init__(parent)
        self.config = config
        self._title_custom = False
        self._folder_custom = False
        self.output_dir: Path | None = None
        self._build_ui()
        self._emit_config()
        self._update_resolution()
        self._update_word_count()

    def _build_ui(self) -> None:
        root = QtWidgets.QHBoxLayout(self)

        # left - scrollable form
        scroll = QtWidgets.QScrollArea()
        scroll.setWidgetResizable(True)
        left = QtWidgets.QWidget()
        form = QtWidgets.QFormLayout(left)

        self.title_edit = QtWidgets.QLineEdit()
        self.title_edit.setToolTip("Project title")
        form.addRow("Title", self.title_edit)

        self.folder_edit = QtWidgets.QLineEdit()
        self.folder_edit.setToolTip("Output folder")
        form.addRow("Folder", self.folder_edit)

        self.script_edit = ScriptEdit()
        self.script_edit.setToolTip("Drop or paste script text")
        form.addRow("Script", self.script_edit)
        self.word_label = QtWidgets.QLabel("0 words")
        self.duration_label = QtWidgets.QLabel("~0 sec")
        word_row = QtWidgets.QHBoxLayout()
        word_row.addWidget(self.word_label)
        word_row.addWidget(self.duration_label)
        word_widget = QtWidgets.QWidget()
        word_widget.setLayout(word_row)
        form.addRow(word_widget)
        self.warning_label = QtWidgets.QLabel()
        form.addRow(self.warning_label)

        self.load_btn = QtWidgets.QPushButton("Load Script")
        self.load_btn.clicked.connect(self._load_script)
        form.addRow(self.load_btn)
        self.download_btn = QtWidgets.QPushButton("Download Background")
        self.download_btn.clicked.connect(self._download_background)
        form.addRow(self.download_btn)

        self.ai_btn = QtWidgets.QPushButton("Generate Story with AI")
        self.ai_btn.clicked.connect(self._ai_popup)
        form.addRow(self.ai_btn)

        snap_row = QtWidgets.QHBoxLayout()
        self.save_btn = QtWidgets.QPushButton("Save Project")
        self.save_btn.clicked.connect(self._save_project)
        self.load_btn2 = QtWidgets.QPushButton("Load Project")
        self.load_btn2.clicked.connect(self._load_project)
        snap_row.addWidget(self.save_btn)
        snap_row.addWidget(self.load_btn2)
        snap_widget = QtWidgets.QWidget()
        snap_widget.setLayout(snap_row)
        form.addRow("Session", snap_widget)

        preset_row = QtWidgets.QHBoxLayout()
        self.preset_save = QtWidgets.QPushButton("Save Preset")
        self.preset_save.clicked.connect(self._save_preset)
        self.preset_load = QtWidgets.QPushButton("Load Preset")
        self.preset_load.clicked.connect(self._load_preset)
        preset_row.addWidget(self.preset_save)
        preset_row.addWidget(self.preset_load)
        preset_widget = QtWidgets.QWidget()
        preset_widget.setLayout(preset_row)
        form.addRow("Export Preset", preset_widget)

        self.voice_combo = QtWidgets.QComboBox()
        for name, vid in (self.config.voices or {}).items():
            self.voice_combo.addItem(name, vid)
        self.voice_combo.setToolTip("Select voice")
        voice_row = QtWidgets.QHBoxLayout()
        voice_row.addWidget(self.voice_combo, 1)
        self.preview_btn = QtWidgets.QPushButton("\u25B6\ufe0f Preview")
        self.preview_btn.clicked.connect(self._preview_voice)
        voice_row.addWidget(self.preview_btn)
        voice_widget = QtWidgets.QWidget()
        voice_widget.setLayout(voice_row)
        form.addRow("Voice", voice_widget)

        self.style_combo = QtWidgets.QComboBox()
        self.style_combo.addItems(["simple", "karaoke", "progressive"])
        self.style_combo.setToolTip("Subtitle style")
        form.addRow("Subtitle Style", self.style_combo)

        self.bg_combo = QtWidgets.QComboBox()
        for name in (self.config.background_styles or {}).keys():
            self.bg_combo.addItem(name)
        self.bg_combo.setToolTip("Background style")
        form.addRow("Background", self.bg_combo)

        self.platform_combo = QtWidgets.QComboBox()
        self.platform_combo.addItems(["TikTok", "Reels", "Shorts", "Square"])
        self.resolution_label = QtWidgets.QLabel()
        platform_row = QtWidgets.QHBoxLayout()
        platform_row.addWidget(self.platform_combo)
        platform_row.addWidget(self.resolution_label)
        pwidget = QtWidgets.QWidget()
        pwidget.setLayout(platform_row)
        self.platform_combo.setToolTip("Target platform")
        form.addRow("Platform", pwidget)

        self.adv_box = QtWidgets.QGroupBox("Advanced Settings")
        self.adv_box.setCheckable(True)
        self.adv_box.setChecked(False)
        adv_layout = QtWidgets.QFormLayout(self.adv_box)

        self.format_combo = QtWidgets.QComboBox()
        self.format_combo.addItems(["mp4", "webm", "mov"])
        adv_layout.addRow("Format", self.format_combo)

        self.fps_combo = QtWidgets.QComboBox()
        self.fps_combo.addItems(["24", "30", "60"])
        adv_layout.addRow("Frame Rate", self.fps_combo)

        self.watermark_check = QtWidgets.QCheckBox("Include Watermark")
        self.watermark_check.setChecked(True)
        self.watermark_check.setToolTip("Add watermark overlay")
        adv_layout.addRow(self.watermark_check)

        self.trim_check = QtWidgets.QCheckBox("Auto-Trim Silence")
        self.trim_check.setToolTip("Remove silence from voiceover")
        adv_layout.addRow(self.trim_check)
        self.crop_check = QtWidgets.QCheckBox("Crop Safe Zone")
        self.crop_check.setToolTip("Crop video to safe area")
        adv_layout.addRow(self.crop_check)
        self.summary_check = QtWidgets.QCheckBox("Summary Overlay")
        self.summary_check.setToolTip("Add summary card on video")
        adv_layout.addRow(self.summary_check)
        form.addRow(self.adv_box)
        self.adv_box.toggled.connect(self._toggle_advanced)
        self._toggle_advanced(False)

        # intro/outro upload
        self.intro_check = QtWidgets.QCheckBox("Use Intro")
        intro_row = QtWidgets.QHBoxLayout()
        self.intro_path = QtWidgets.QLineEdit()
        self.intro_browse = QtWidgets.QPushButton("Browse")
        self.intro_browse.clicked.connect(lambda: self._browse_file(self.intro_path))
        intro_row.addWidget(self.intro_check)
        intro_row.addWidget(self.intro_path, 1)
        intro_row.addWidget(self.intro_browse)
        intro_widget = QtWidgets.QWidget()
        intro_widget.setLayout(intro_row)
        form.addRow("Intro Video", intro_widget)

        self.outro_check = QtWidgets.QCheckBox("Use Outro")
        outro_row = QtWidgets.QHBoxLayout()
        self.outro_path = QtWidgets.QLineEdit()
        self.outro_browse = QtWidgets.QPushButton("Browse")
        self.outro_browse.clicked.connect(lambda: self._browse_file(self.outro_path))
        outro_row.addWidget(self.outro_check)
        outro_row.addWidget(self.outro_path, 1)
        outro_row.addWidget(self.outro_browse)
        outro_widget = QtWidgets.QWidget()
        outro_widget.setLayout(outro_row)
        form.addRow("Outro Video", outro_widget)

        # subtitle style editor
        style_box = QtWidgets.QGroupBox("Subtitle Style Editor")
        style_layout = QtWidgets.QFormLayout(style_box)
        self.font_combo = QtWidgets.QFontComboBox()
        style_layout.addRow("Font", self.font_combo)
        self.font_size = QtWidgets.QSpinBox()
        self.font_size.setRange(10, 100)
        self.font_size.setValue(48)
        style_layout.addRow("Size", self.font_size)
        self.color_btn = QtWidgets.QPushButton("Color")
        self.color_btn.clicked.connect(self._pick_color)
        self.color_btn.setToolTip("Subtitle color")
        style_layout.addRow(self.color_btn)
        self.outline_check = QtWidgets.QCheckBox("Outline/Shadow")
        self.outline_check.setChecked(True)
        style_layout.addRow(self.outline_check)
        self.position_slider = QtWidgets.QSlider(QtCore.Qt.Orientation.Horizontal)
        self.position_slider.setRange(40, 80)
        self.position_slider.setValue(55)
        style_layout.addRow("Height %", self.position_slider)
        self.sample_label = QtWidgets.QLabel("Sample subtitle")
        self.sample_label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        style_layout.addRow(self.sample_label)
        form.addRow(style_box)

        self.generate_btn = QtWidgets.QPushButton("Generate Final Video")
        self.generate_btn.clicked.connect(self._emit_generate)
        form.addRow(self.generate_btn)

        self._connect_signals()

        self.output_label = QtWidgets.QLabel()
        self.output_label.setOpenExternalLinks(True)
        form.addRow("Output", self.output_label)
        status_row = QtWidgets.QHBoxLayout()
        self.status_led = QtWidgets.QLabel()
        self.status_led.setFixedSize(12, 12)
        self.status_led.setStyleSheet("background:#555;border-radius:6px")
        self.status_label = QtWidgets.QLabel("Idle")
        status_row.addWidget(self.status_led)
        status_row.addWidget(self.status_label)
        status_widget = QtWidgets.QWidget()
        status_widget.setLayout(status_row)
        form.addRow("Status", status_widget)

        self.log_edit = QtWidgets.QPlainTextEdit()
        self.log_edit.setReadOnly(True)
        self.log_edit.setMaximumHeight(100)
        form.addRow(self.log_edit)

        self.dev_check = QtWidgets.QCheckBox("Developer Mode")
        self.dev_check.setChecked(self.config.developer_mode)
        form.addRow(self.dev_check)
        self.json_view = QtWidgets.QPlainTextEdit()
        self.json_view.setReadOnly(True)
        self.json_view.setVisible(False)
        form.addRow(self.json_view)

        scroll.setWidget(left)
        root.addWidget(scroll, 1)

        # right preview panel
        right = QtWidgets.QVBoxLayout()

        self.preview_box = QtWidgets.QFrame()
        self.preview_box.setObjectName("PreviewPhone")
        self.preview_box.setFixedSize(270, 480)
        stack = QtWidgets.QStackedLayout(self.preview_box)
        self.video_widget = QtMultimediaWidgets.QVideoWidget()
        self.player = QtMultimedia.QMediaPlayer()
        self.player.setVideoOutput(self.video_widget)
        stack.addWidget(self.video_widget)
        self.thumbnail = QtWidgets.QLabel()
        self.thumbnail.setScaledContents(True)
        self.thumbnail.setVisible(False)
        stack.addWidget(self.thumbnail)
        self.subtitle_preview = QtWidgets.QLabel("Test subtitle")
        self.subtitle_preview.setStyleSheet("color:white; background: transparent;")
        self.subtitle_preview.setAlignment(QtCore.Qt.AlignBottom | QtCore.Qt.AlignHCenter)
        stack.addWidget(self.subtitle_preview)
        stack.setStackingMode(QtWidgets.QStackedLayout.StackAll)
        self._start_sample_video()
        right.addWidget(self.preview_box, alignment=QtCore.Qt.AlignmentFlag.AlignCenter)
        self.preview_summary = QtWidgets.QLabel()
        self.preview_summary.setWordWrap(True)
        right.addWidget(self.preview_summary)
        btn_row = QtWidgets.QHBoxLayout()
        self.open_btn = QtWidgets.QPushButton("Open Folder")
        self.open_btn.clicked.connect(self._open_folder)
        btn_row.addWidget(self.open_btn)
        self.copy_btn = QtWidgets.QPushButton("Copy for TikTok")
        self.copy_btn.clicked.connect(self._copy_path)
        btn_row.addWidget(self.copy_btn)
        self.sub_btn = QtWidgets.QPushButton("Preview Subtitles")
        self.sub_btn.clicked.connect(self._preview_subs)
        btn_row.addWidget(self.sub_btn)
        right.addLayout(btn_row)

        right_widget = QtWidgets.QWidget()
        right_widget.setLayout(right)
        root.addWidget(right_widget)

    def _connect_signals(self) -> None:
        widgets = [
            self.title_edit,
            self.folder_edit,
            self.script_edit,
            self.voice_combo,
            self.style_combo,
            self.bg_combo,
            self.platform_combo,
            self.format_combo,
            self.fps_combo,
            self.watermark_check,
            self.trim_check,
            self.crop_check,
            self.summary_check,
            self.intro_check,
            self.font_combo,
            self.font_size,
            self.outline_check,
            self.position_slider,
            self.adv_box,
            self.dev_check,
            self.preset_save,
            self.preset_load,
        ]
        for widget in widgets:
            if isinstance(widget, QtWidgets.QComboBox):
                widget.currentIndexChanged.connect(self._emit_config)
            elif isinstance(widget, (QtWidgets.QCheckBox, QtWidgets.QPushButton)):
                widget.toggled.connect(self._emit_config) if hasattr(widget, 'toggled') else widget.clicked.connect(self._emit_config)
            elif isinstance(widget, QtWidgets.QSlider):
                widget.valueChanged.connect(self._emit_config)
            else:
                widget.textChanged.connect(self._emit_config)
        self.script_edit.textChanged.connect(self._update_word_count)
        if isinstance(self.script_edit, ScriptEdit):
            self.script_edit.fileDropped.connect(self._load_dropped_file)
        self.title_edit.textEdited.connect(self._flag_title_custom)
        self.folder_edit.textEdited.connect(self._flag_folder_custom)
        self.platform_combo.currentIndexChanged.connect(self._update_resolution)

    def _load_script(self) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Select Script", "scripts", "Text Files (*.txt *.docx)")
        if not path:
            return
        p = Path(path)
        if p.suffix.lower() == ".docx":
            try:
                import docx  # type: ignore
            except Exception:
                QtWidgets.QMessageBox.warning(self, "Load", "python-docx not installed")
                return
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            text = p.read_text()
        self.script_edit.setPlainText(text)
        self.title_edit.setText(p.stem)
        self._title_custom = False
        self._folder_custom = False
        self._update_word_count()
        self._suggest_title()

    def _load_dropped_file(self, path: str) -> None:
        p = Path(path)
        if not p.exists():
            return
        if p.suffix.lower() == ".docx":
            try:
                import docx  # type: ignore
            except Exception:
                return
            doc = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            text = p.read_text()
        self.script_edit.setPlainText(text)
        self.title_edit.setText(p.stem)
        self._title_custom = False
        self._folder_custom = False
        self._update_word_count()
        self._suggest_title()

    def _flag_title_custom(self) -> None:
        self._title_custom = True

    def _flag_folder_custom(self) -> None:
        self._folder_custom = True

    def _toggle_advanced(self, state: bool) -> None:
        for child in self.adv_box.findChildren(QtWidgets.QWidget):
            if child is not self.adv_box:
                child.setVisible(state)
        self._emit_config()

    def _ai_popup(self) -> None:
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Generate with AI")
        form = QtWidgets.QFormLayout(dlg)
        genre = QtWidgets.QComboBox()
        genre.addItems(["Creepypasta", "Urban Legend", "Horror", "Mystery"])
        templates = {
            "Urban Legend": "Write a scary urban legend about a haunted city bus.",
            "True Crime": "Summarize a true crime story in a suspenseful tone.",
            "Mystery": "Craft a mysterious short story involving an abandoned house.",
        }
        template_combo = QtWidgets.QComboBox()
        template_combo.addItem("None")
        template_combo.addItems(list(templates.keys()))
        prompt_edit = QtWidgets.QPlainTextEdit()
        form.addRow("Genre", genre)
        form.addRow("Template", template_combo)
        form.addRow("Prompt", prompt_edit)
        def set_template(idx: int) -> None:
            name = template_combo.currentText()
            if name != "None":
                prompt_edit.setPlainText(templates.get(name, ""))
        template_combo.currentIndexChanged.connect(set_template)
        btns = QtWidgets.QDialogButtonBox(
            QtWidgets.QDialogButtonBox.Ok | QtWidgets.QDialogButtonBox.Cancel
        )
        form.addRow(btns)
        btns.accepted.connect(dlg.accept)
        btns.rejected.connect(dlg.reject)
        if dlg.exec() == QtWidgets.QDialog.Accepted:
            prompt = prompt_edit.toPlainText().strip()
            g = genre.currentText()
            text = generator.generate_story(g, None, prompt)
            self.script_edit.setPlainText(text)
            self.log_edit.appendPlainText("[AI] Story generated")
            self._update_word_count()
            self._title_custom = False
            self._suggest_title()

    def _emit_generate(self) -> None:
        self.generateRequested.emit(self._collect_opts())

    def _collect_opts(self) -> dict:
        return {
            "title": self.title_edit.text().strip(),
            "script": self.script_edit.toPlainText().strip(),
            "folder": self.folder_edit.text().strip(),
            "voice": self.voice_combo.currentData(),
            "subtitle_style": self.style_combo.currentText(),
            "background": self.bg_combo.currentText(),
            "platform": self.platform_combo.currentText(),
            "format": self.format_combo.currentText(),
            "fps": int(self.fps_combo.currentText()),
            "watermark": self.watermark_check.isChecked(),
            "trim_silence": self.trim_check.isChecked(),
            "crop_safe": self.crop_check.isChecked(),
            "summary_overlay": self.summary_check.isChecked(),
            "intro": self.intro_check.isChecked(),
            "intro_path": self.intro_path.text().strip(),
            "outro": self.outro_check.isChecked(),
            "outro_path": self.outro_path.text().strip(),
            "font": self.font_combo.currentFont().family(),
            "font_size": self.font_size.value(),
            "color": self.color_btn.palette().button().color().name(),
            "outline": self.outline_check.isChecked(),
            "height": self.position_slider.value(),
            "developer_mode": self.dev_check.isChecked(),
        }

    def _emit_config(self) -> None:
        opts = self._collect_opts()
        self.configChanged.emit(opts)
        if self.dev_check.isChecked():
            import json

            self.json_view.setPlainText(json.dumps(opts, indent=2))
            self.json_view.setVisible(True)
        else:
            self.json_view.setVisible(False)

    def _preview_voice(self) -> None:
        preview_voice(
            self.config.voice_engine,
            self.voice_combo.currentData(),
            self.config.coqui_model_name,
        )

    def _download_background(self) -> None:
        url, ok = QtWidgets.QInputDialog.getText(self, "Download Background", "YouTube URL")
        if not ok or not url:
            return
        dest = Path("assets/backgrounds/custom")
        dest.mkdir(parents=True, exist_ok=True)
        tmp = dest / "temp.mp4"
        try:
            from yt_dlp import YoutubeDL

            ydl = YoutubeDL({"outtmpl": str(tmp)})
            ydl.download([url])
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Download", f"Failed: {e}")
            self.log_edit.appendPlainText(f"[DL] failed {e}")
            return
        duration = self._estimate_duration()
        out_path = dest / f"bg_{datetime.now().strftime('%H%M%S')}.mp4"
        cmd = [
            "ffmpeg",
            "-y",
            "-i",
            str(tmp),
            "-t",
            str(int(duration)),
            str(out_path),
        ]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            tmp.unlink(missing_ok=True)
            if self.bg_combo.findText("Custom") == -1:
                self.bg_combo.addItem("Custom")
            self.log_edit.appendPlainText(f"[DL] saved {out_path}")
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "FFmpeg", f"Failed trim: {e}")
            self.log_edit.appendPlainText(f"[DL] ffmpeg fail {e}")

    def _update_word_count(self) -> None:
        words = len(self.script_edit.toPlainText().split())
        self.word_label.setText(f"{words} words")
        duration = self._estimate_duration()
        self.duration_label.setText(f"~{int(duration)} sec")
        warn = ""
        if self.platform_combo.currentText() in {"TikTok", "Reels", "Shorts"}:
            if duration > 60:
                warn = "Too long for platform"
            elif duration < 5:
                warn = "Too short"
        self.warning_label.setText(warn)
        self._suggest_title()
        self._update_output_preview()

    def _update_resolution(self) -> None:
        platform = self.platform_combo.currentText()
        res = "1080x1920"
        if platform == "Square":
            res = "1080x1080"
        self.resolution_label.setText(res)
        self._update_output_preview()

    def set_status(self, state: str, message: str = "") -> None:
        colors = {
            "idle": "#555",
            "running": "#ffaa00",
            "error": "#d63c3c",
            "success": "#2ecc71",
        }
        self.status_led.setStyleSheet(
            f"background:{colors.get(state, '#555')};border-radius:6px"
        )
        self.status_label.setText(message or state.capitalize())

    def _estimate_duration(self) -> float:
        words = len(self.script_edit.toPlainText().split())
        return words * 0.5

    def _suggest_title(self) -> None:
        if self._title_custom:
            return
        text = self.script_edit.toPlainText().strip()
        if not text:
            return
        first = text.splitlines()[0]
        base = sanitize_name("_".join(first.lower().split()[:3]))
        date = datetime.now().strftime("%Y%m%d")
        self.title_edit.setText(f"{base}_{date}")
        self._update_output_preview()

    def _update_output_preview(self) -> None:
        title = self.title_edit.text().strip() or "session"
        if not self._folder_custom:
            folder = f"output/{sanitize_name(title)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            self.folder_edit.setText(folder)

    # ----- preview helpers -----
    def set_thumbnail(self, path: Path) -> None:
        pix = QtGui.QPixmap(str(path))
        self.thumbnail.setPixmap(pix)
        self.thumbnail.setVisible(True)
        self.player.pause()

    def set_preview_summary(self, text: str) -> None:
        self.preview_summary.setText(text)

    def set_output(self, folder: Path) -> None:
        self.output_dir = folder
        self.output_label.setText(f"<a href='{folder}'>{folder}</a>")
        self.output_label.setOpenExternalLinks(True)

    def _open_folder(self) -> None:
        if not self.output_dir:
            return
        QtGui.QDesktopServices.openUrl(QtCore.QUrl.fromLocalFile(str(self.output_dir)))

    def _copy_path(self) -> None:
        if not self.output_dir:
            return
        cb = QtWidgets.QApplication.clipboard()
        cb.setText(str(self.output_dir / 'final_video.mp4'))

    def _preview_subs(self) -> None:
        if not self.output_dir:
            return
        path = self.output_dir / 'subtitles.ass'
        QtGui.QDesktopServices.openUrl(QtCore.QUrl.fromLocalFile(str(path)))

    def _pick_color(self) -> None:
        color = QtWidgets.QColorDialog.getColor(QtGui.QColor('white'), self)
        if color.isValid():
            self.color_btn.setStyleSheet(f'background:{color.name()}')
            self.sample_label.setStyleSheet(f'color:{color.name()}')
            self._emit_config()

    # ----- session helpers -----
    def _browse_file(self, line_edit: QtWidgets.QLineEdit) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Select Video", "", "Video Files (*.mp4 *.webm *.mov)")
        if path:
            line_edit.setText(path)

    def _save_project(self) -> None:
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Project", "project.acproject", "AC Project (*.acproject)")
        if not path:
            return
        data = {
            "opts": self._collect_opts(),
            "script": self.script_edit.toPlainText(),
        }
        Path(path).write_text(json.dumps(data, indent=2))
        self.log_edit.appendPlainText(f"[Session] Saved {path}")

    def _load_project(self) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Load Project", "", "AC Project (*.acproject)")
        if not path:
            return
        try:
            data = json.loads(Path(path).read_text())
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Load", f"Failed: {e}")
            return
        opts = data.get("opts", {})
        self.script_edit.setPlainText(data.get("script", ""))
        self._apply_opts(opts)
        self.log_edit.appendPlainText(f"[Session] Loaded {path}")

    def _save_preset(self) -> None:
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Preset", "export.json", "JSON (*.json)")
        if not path:
            return
        data = {
            "format": self.format_combo.currentText(),
            "fps": self.fps_combo.currentText(),
            "watermark": self.watermark_check.isChecked(),
            "font_size": self.font_size.value(),
            "color": self.color_btn.palette().button().color().name(),
        }
        Path(path).write_text(json.dumps(data, indent=2))
        self.log_edit.appendPlainText(f"[Preset] Saved {path}")

    def _load_preset(self) -> None:
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Load Preset", "", "JSON (*.json)")
        if not path:
            return
        try:
            data = json.loads(Path(path).read_text())
        except Exception as e:
            QtWidgets.QMessageBox.warning(self, "Load", f"Failed: {e}")
            return
        self.format_combo.setCurrentText(data.get("format", self.format_combo.currentText()))
        self.fps_combo.setCurrentText(str(data.get("fps", self.fps_combo.currentText())))
        self.watermark_check.setChecked(data.get("watermark", True))
        self.font_size.setValue(int(data.get("font_size", self.font_size.value())))
        color = data.get("color")
        if color:
            self.color_btn.setStyleSheet(f"background:{color}")
            self.sample_label.setStyleSheet(f"color:{color}")
        self.log_edit.appendPlainText(f"[Preset] Loaded {path}")

    def _apply_opts(self, opts: dict) -> None:
        self.title_edit.setText(opts.get("title", ""))
        self.folder_edit.setText(opts.get("folder", ""))
        self.voice_combo.setCurrentIndex(max(0, self.voice_combo.findData(opts.get("voice"))))
        self.style_combo.setCurrentText(opts.get("subtitle_style", self.style_combo.currentText()))
        self.bg_combo.setCurrentText(opts.get("background", self.bg_combo.currentText()))
        self.platform_combo.setCurrentText(opts.get("platform", self.platform_combo.currentText()))
        self.format_combo.setCurrentText(opts.get("format", self.format_combo.currentText()))
        self.fps_combo.setCurrentText(str(opts.get("fps", self.fps_combo.currentText())))
        self.watermark_check.setChecked(opts.get("watermark", True))
        self.trim_check.setChecked(opts.get("trim_silence", False))
        self.crop_check.setChecked(opts.get("crop_safe", False))
        self.summary_check.setChecked(opts.get("summary_overlay", False))
        self.intro_check.setChecked(opts.get("intro", False))
        self.intro_path.setText(opts.get("intro_path", ""))
        self.outro_check.setChecked(opts.get("outro", False))
        self.outro_path.setText(opts.get("outro_path", ""))
        self.font_combo.setCurrentFont(QtGui.QFont(opts.get("font", self.font_combo.currentFont().family())))
        self.font_size.setValue(int(opts.get("font_size", self.font_size.value())))
        if "color" in opts:
            self.color_btn.setStyleSheet(f"background:{opts['color']}")
            self.sample_label.setStyleSheet(f"color:{opts['color']}")
        self.outline_check.setChecked(opts.get("outline", True))
        self.position_slider.setValue(int(opts.get("height", self.position_slider.value())))
        self.dev_check.setChecked(opts.get("developer_mode", self.dev_check.isChecked()))
        self._title_custom = True
        self._folder_custom = True
        self._update_word_count()
        self._emit_config()

    def reset_form(self) -> None:
        self.title_edit.clear()
        self.folder_edit.clear()
        self.script_edit.clear()
        self._title_custom = False
        self._folder_custom = False
        self._update_word_count()
        self._suggest_title()

    # ----- error inspector -----
    def show_error(self, trace: str) -> None:
        if not hasattr(self, "error_view"):
            self.error_view = QtWidgets.QPlainTextEdit()
            self.error_view.setReadOnly(True)
            self.layout().itemAt(0).widget().layout().addRow("Error", self.error_view)
        self.error_view.setPlainText(trace)
        self.error_view.setVisible(bool(trace))

    def _start_sample_video(self) -> None:
        paths = list(Path(self.config.background_videos_path).rglob("*.mp4"))
        if not paths:
            return
        self.player.setSource(QtCore.QUrl.fromLocalFile(str(paths[0])))
        try:
            self.player.setLoops(QtMultimedia.QMediaPlayer.Loops.Infinite)
        except Exception:
            try:
                self.player.setLoops(QtMultimedia.QMediaPlayer.Infinite)
            except Exception:
                pass
        self.player.play()
