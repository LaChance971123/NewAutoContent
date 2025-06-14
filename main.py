import sys
import json
import random
from pathlib import Path

GUI_DIR = Path(__file__).parent / "PyOneDark_GUI_Core"
sys.path.insert(0, str(GUI_DIR))

from qt_core import *
from gui.core.json_settings import Settings
from gui.core.json_themes import Themes
from gui.widgets import PyPushButton, PyToggle, PyGrips
from pipeline import generator
from pipeline.pipeline import VideoPipeline
from pipeline.config import Config
from pipeline.helpers import sanitize_name
from docx import Document

DEFAULTS = {
    "api_key": "",
    "output_folder": "output",
    "use_coqui_fallback": False,
    "subtitle_font": "Noto Sans",
    "output_resolution": "1080x1920",
    "watermark": False,
    "ai_prompt": "",
    "max_story_len": 200,
    "last_voice": "Default",
}

SECTION_STYLE = (
    "font-weight: bold; font-size: 12pt; margin-bottom: 12px;"
)
PAD = 10


class ScriptEdit(QPlainTextEdit):
    """Editor that supports drag-and-drop for text and docx files."""

    fileLoaded = Signal(str)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setAcceptDrops(True)

    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls():
            url = e.mimeData().urls()[0]
            if url.isLocalFile() and url.toLocalFile().lower().endswith((".txt", ".docx")):
                e.acceptProposedAction()
                return
        super().dragEnterEvent(e)

    def dropEvent(self, e):
        if e.mimeData().hasUrls():
            path = e.mimeData().urls()[0].toLocalFile()
            lower = path.lower()
            try:
                if lower.endswith(".txt"):
                    with open(path, "r", encoding="utf-8") as f:
                        self.setPlainText(f.read())
                    self.fileLoaded.emit(Path(path).name)
                    e.acceptProposedAction()
                    return
                elif lower.endswith(".docx"):
                    doc = Document(path)
                    text = "\n".join(p.text for p in doc.paragraphs)
                    self.setPlainText(text)
                    self.fileLoaded.emit(Path(path).name)
                    e.acceptProposedAction()
                    return
                else:
                    QMessageBox.warning(self, "Unsupported", "Only .txt and .docx files are supported")
                    e.ignore()
                    return
            except Exception as exc:
                QMessageBox.warning(self, "Error", str(exc))
                e.ignore()
                return
        super().dropEvent(e)


class SettingsManager:
    """Simple wrapper around the global settings.json"""

    def __init__(self):
        self.path = Path(Settings.settings_path)
        self.load()

    def load(self):
        try:
            self.data = json.loads(self.path.read_text())
        except Exception:
            self.data = {}
        for k, v in DEFAULTS.items():
            self.data.setdefault(k, v)

    def save(self):
        with open(self.path, "w", encoding="utf-8") as f:
            json.dump(self.data, f, indent=4)


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.settings = Settings().items
        Themes.settings_path = str(GUI_DIR / f"gui/themes/{self.settings['theme_name']}.json")
        self.themes = Themes().items
        self.app_settings = SettingsManager()
        self.auto_filename = True

        self.setup_ui()
        self.build_sidebar()
        self.build_home_page()
        self.build_settings_page()
        self.build_help_page()
        self.status_timer = QTimer(self)
        self.status_timer.setSingleShot(True)
        self.status_timer.timeout.connect(lambda: self.set_status("Idle", self.themes["app_color"].get("green")))
        self.reset_form()
        self.restore_state()
        self.show()

    # ------------------------------------------------------------------
    def setup_ui(self):
        from gui.uis.windows.main_window.ui_main import UI_MainWindow

        self.ui = UI_MainWindow()
        self.ui.setup_ui(self)

        if self.settings["custom_title_bar"]:
            self.setWindowFlag(Qt.FramelessWindowHint)
            self.setAttribute(Qt.WA_TranslucentBackground)
        self.setWindowTitle(self.settings["app_name"])

        if self.settings["custom_title_bar"]:
            self.left_grip = PyGrips(self, "left", True)
            self.right_grip = PyGrips(self, "right", True)
            self.top_grip = PyGrips(self, "top", True)
            self.bottom_grip = PyGrips(self, "bottom", True)
            self.top_left_grip = PyGrips(self, "top_left", True)
            self.top_right_grip = PyGrips(self, "top_right", True)
            self.bottom_left_grip = PyGrips(self, "bottom_left", True)
            self.bottom_right_grip = PyGrips(self, "bottom_right", True)

        # clear default left menu
        lm = self.ui.left_menu_frame.layout()
        while lm.count():
            item = lm.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self.menu_container = QFrame()
        self.sidebar_layout = QVBoxLayout(self.menu_container)
        self.sidebar_layout.setContentsMargins(PAD, PAD, PAD, PAD)
        lm.addWidget(self.menu_container)

        # hide left/right columns by default
        self.ui.left_column_frame.setMinimumWidth(0)
        self.ui.left_column_frame.setMaximumWidth(0)
        self.ui.right_column_frame.setMinimumWidth(0)
        self.ui.right_column_frame.setMaximumWidth(0)

    # ------------------------------------------------------------------
    def build_sidebar(self):
        theme = self.themes["app_color"]
        btn_args = dict(
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )

        self.home_btn = PyPushButton(text="Home", **btn_args)
        self.settings_btn = PyPushButton(text="Settings", **btn_args)
        self.help_btn = PyPushButton(text="Help", **btn_args)

        self.home_btn.clicked.connect(lambda: self.switch_page(0))
        self.settings_btn.clicked.connect(lambda: self.switch_page(1))
        self.help_btn.clicked.connect(lambda: self.switch_page(2))

        self.sidebar_layout.addWidget(self.home_btn)
        self.sidebar_layout.addWidget(self.settings_btn)
        self.sidebar_layout.addWidget(self.help_btn)
        self.sidebar_layout.addStretch()

    # ------------------------------------------------------------------
    def switch_page(self, index: int):
        stack = self.ui.load_pages.pages
        if index == stack.currentIndex():
            return
        new = stack.widget(index)
        effect = QGraphicsOpacityEffect(new)
        new.setGraphicsEffect(effect)
        anim = QPropertyAnimation(effect, b"opacity", self)
        anim.setDuration(200)
        anim.setStartValue(0)
        anim.setEndValue(1)
        anim.start(QPropertyAnimation.DeleteWhenStopped)
        stack.setCurrentIndex(index)
        if hasattr(self, "preview_container"):
            self.preview_container.setVisible(index == 0)

    # ------------------------------------------------------------------
    def build_home_page(self):
        theme = self.themes["app_color"]
        layout = self.ui.load_pages.controls_layout
        layout.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        layout.setContentsMargins(PAD, PAD, PAD, PAD)
        self.ui.load_pages.preview_layout.setContentsMargins(PAD, PAD, PAD, PAD)

        def section(title: str):
            frame = QFrame()
            frame.setMaximumWidth(520)
            fl = QVBoxLayout(frame)
            fl.setContentsMargins(PAD, PAD, PAD, PAD)
            fl.setSpacing(6)
            lbl = QLabel(title)
            lbl.setObjectName("section")
            lbl.setStyleSheet(SECTION_STYLE)
            fl.addWidget(lbl)
            layout.addWidget(frame, alignment=Qt.AlignHCenter)
            return fl

        # "🎬 Script Input" section with title and script boxes
        title_layout = section("\U0001F3AC Script Input")
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("Title / Output Filename")
        self.title_edit.setMaximumWidth(500)
        self.title_edit.textEdited.connect(self.disable_auto_title)
        title_layout.addWidget(self.title_edit)
        script_layout = title_layout
        self.script_edit = ScriptEdit()
        self.script_edit.setPlaceholderText(
            "Drop your story here or click 'Generate with AI' to begin..."
        )
        self.script_edit.setMinimumHeight(120)
        self.script_edit.setMaximumWidth(500)
        self.script_edit.setStyleSheet(
            f"background-color: {theme['dark_one']}; color: {theme['text_foreground']};"
        )
        self.script_edit.setToolTip("Drag .txt or .docx files here")
        self.script_edit.textChanged.connect(self.on_script_changed)
        self.script_edit.fileLoaded.connect(self.on_script_loaded)
        script_layout.addWidget(self.script_edit)

        btn_row = QHBoxLayout()
        self.upload_btn = PyPushButton(
            text="Upload",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.upload_btn.clicked.connect(self.upload_script)
        self.upload_btn.setMaximumWidth(500)
        btn_row.addWidget(self.upload_btn)

        self.ai_btn = PyPushButton(
            text="Generate with AI",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.ai_btn.setToolTip("Generate a horror story using the local model")
        self.ai_btn.clicked.connect(self.generate_ai_story)
        self.ai_btn.setMaximumWidth(500)
        btn_row.addWidget(self.ai_btn)

        self.reset_btn = PyPushButton(
            text="Reset",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.reset_btn.setToolTip("Clear script")
        self.reset_btn.clicked.connect(self.reset_form)
        self.reset_btn.setMaximumWidth(500)
        btn_row.addWidget(self.reset_btn)
        script_layout.addLayout(btn_row)

        voice_layout = section("\U0001F399\ufe0f Voice Settings")
        row1 = QHBoxLayout()
        self.voice_combo = QComboBox()
        voices = (self.settings.get("voices") or {}).keys()
        self.voice_combo.addItems(list(voices) or ["Default"])
        self.voice_combo.setMaximumWidth(500)
        self.voice_combo.currentTextChanged.connect(lambda t: self.update_setting("last_voice", t))
        row1.addWidget(self.voice_combo)
        self.preview_btn = PyPushButton(
            text="Preview",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.preview_btn.setToolTip("Preview selected voice")
        self.preview_btn.setMaximumWidth(500)
        self.preview_btn.clicked.connect(self.preview_voice)
        row1.addWidget(self.preview_btn)
        voice_layout.addLayout(row1)

        output_layout = section("\u2699\ufe0f Output Options")
        self.subtitle_combo = QComboBox()
        self.subtitle_combo.addItems(["karaoke", "progressive", "simple"])
        self.subtitle_combo.setToolTip("Subtitle style")
        self.subtitle_combo.setMaximumWidth(500)
        self.subtitle_combo.currentTextChanged.connect(lambda t: self.update_setting("last_subtitle", t))
        output_layout.addWidget(self.subtitle_combo)

        wm_row = QHBoxLayout()
        self.watermark_toggle = PyToggle(
            bg_color=theme["dark_three"],
            circle_color=theme["icon_color"],
            active_color=theme["context_color"],
        )
        self.watermark_toggle.toggled.connect(self.update_watermark_label)
        self.watermark_toggle.toggled.connect(lambda v: self.update_setting("watermark", v))
        wm_row.addWidget(self.watermark_toggle)
        self.watermark_label = QLabel("Watermark: Off")
        wm_row.addWidget(self.watermark_label)
        wm_row.addStretch()
        output_layout.addLayout(wm_row)

        bg_layout = section("\U0001F39E\ufe0f Background Style")
        row2 = QHBoxLayout()
        self.bg_combo = QComboBox()
        bg_styles = (self.settings.get("background_styles") or {}).keys()
        self.bg_combo.addItems(list(bg_styles) or ["Default"])
        self.bg_combo.setMaximumWidth(500)
        self.bg_combo.currentTextChanged.connect(lambda t: self.update_setting("last_background", t))
        row2.addWidget(self.bg_combo)
        self.surprise_btn = PyPushButton(
            text="Surprise Me",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.surprise_btn.setToolTip("Randomly select options")
        self.surprise_btn.setMaximumWidth(500)
        self.surprise_btn.clicked.connect(self.surprise_me)
        row2.addWidget(self.surprise_btn)
        bg_layout.addLayout(row2)

        self.output_info = QLabel("1080p @30fps | Watermark Off")
        self.output_info.setAlignment(Qt.AlignCenter)
        output_layout.addWidget(self.output_info)

        self.create_btn = PyPushButton(
            text="Create Content",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["context_color"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.create_btn.setEnabled(False)
        self.create_btn.setToolTip("Run the full pipeline")
        self.create_btn.setMaximumWidth(500)
        self.create_btn.clicked.connect(self.run_pipeline)
        output_layout.addWidget(self.create_btn)

        status = QHBoxLayout()
        self.status_dot = QLabel("\u25CF")
        self.status_dot.setStyleSheet(f"color: {theme['green']}")
        status.addWidget(self.status_dot)
        self.status_text = QLabel("Idle")
        status.addWidget(self.status_text)
        status.addStretch()
        self.ready_label = QLabel("Ready")
        status.addWidget(self.ready_label, alignment=Qt.AlignRight)
        output_layout.addLayout(status)

        layout.addStretch()

        # preview placeholder
        self.preview_container = QFrame(objectName="preview")
        self.preview_container.setStyleSheet(
            f"#preview {{"
            f"background-color: {theme['dark_one']};"
            f"border-radius: 12px;"
            f"border: 1px solid {theme['dark_four']};"
            f"}}"
        )
        self.preview_container.setToolTip(
            "This is a mockup of how your video will appear on mobile."
        )
        shadow = QGraphicsDropShadowEffect(self.preview_container)
        shadow.setBlurRadius(15)
        shadow.setOffset(0, 0)
        shadow.setColor(QColor(theme.get("dark_two", "#000000")))
        self.preview_container.setGraphicsEffect(shadow)
        preview_layout = QVBoxLayout(self.preview_container)
        preview_layout.setContentsMargins(PAD, PAD, PAD, PAD)
        preview_layout.addStretch()
        preview_label = QLabel("Video Preview")
        preview_label.setAlignment(Qt.AlignCenter)
        preview_layout.addWidget(preview_label, 0, Qt.AlignCenter)
        preview_layout.addStretch()
        self.ui.load_pages.preview_layout.addStretch()
        self.ui.load_pages.preview_layout.addWidget(
            self.preview_container, 0, Qt.AlignCenter
        )
        self.ui.load_pages.preview_layout.addStretch()
        self.adjust_preview_size()

    # ------------------------------------------------------------------
    def build_settings_page(self):
        theme = self.themes["app_color"]
        self.ui.load_pages.title_label.setText("Settings")
        self.ui.load_pages.description_label.setText("Application preferences")

        root = self.ui.load_pages.page_2
        layout = self.ui.load_pages.page_2_layout

        # clear existing widgets
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        scroll = QScrollArea()
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setWidgetResizable(True)
        container = QWidget()
        scroll.setWidget(container)
        vbox = QVBoxLayout(container)
        vbox.setAlignment(Qt.AlignTop)

        def section(title: str) -> QVBoxLayout:
            frame = QFrame()
            frame.setMaximumWidth(520)
            fl = QVBoxLayout(frame)
            fl.setContentsMargins(PAD, PAD, PAD, PAD)
            fl.setSpacing(6)
            lbl = QLabel(title)
            lbl.setObjectName("section")
            lbl.setStyleSheet(SECTION_STYLE)
            lbl.setToolTip(title)
            fl.addWidget(lbl)
            vbox.addWidget(frame, alignment=Qt.AlignHCenter)
            return fl

        # General Settings -------------------------------------------------
        gen_layout = section("\U0001F3A7 General Settings")

        self.resolution_combo = QComboBox()
        self.resolution_combo.addItems(self.settings.get("resolutions", ["1080x1920", "720x1280"]))
        self.resolution_combo.setCurrentText(self.app_settings.data.get("output_resolution", DEFAULTS["output_resolution"]))
        self.resolution_combo.currentTextChanged.connect(lambda t: self.update_setting("output_resolution", t))
        self.resolution_combo.setToolTip("Output resolution")
        gen_layout.addWidget(self.resolution_combo)

        wm_row = QHBoxLayout()
        self.wm_check = QCheckBox("Enable watermark")
        self.wm_check.setChecked(self.app_settings.data.get("watermark", DEFAULTS["watermark"]))
        self.wm_check.toggled.connect(lambda v: self.update_setting("watermark", v))
        self.wm_check.setToolTip("Toggle watermark on output video")
        wm_row.addWidget(self.wm_check)
        wm_row.addStretch()
        gen_layout.addLayout(wm_row)

        out_row = QHBoxLayout()
        self.output_edit = QLineEdit(self.app_settings.data.get("output_folder", DEFAULTS["output_folder"]))
        self.output_edit.setReadOnly(True)
        self.output_edit.setToolTip("Folder where videos are saved")
        out_row.addWidget(self.output_edit)
        self.output_btn = PyPushButton(
            text="Browse",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.output_btn.clicked.connect(self.choose_output_folder)
        out_row.addWidget(self.output_btn)
        gen_layout.addLayout(out_row)

        # Voice Settings ---------------------------------------------------
        voice_layout = section("\U0001F399\ufe0f Voice Settings")

        self.default_voice_combo = QComboBox()
        voices = (self.settings.get("voices") or {}).keys()
        self.default_voice_combo.addItems(list(voices) or ["Default"])
        self.default_voice_combo.setCurrentText(self.app_settings.data.get("last_voice", DEFAULTS["last_voice"]))
        self.default_voice_combo.currentTextChanged.connect(lambda t: (self.update_setting("last_voice", t), self.sync_voice_combo(t)))
        self.default_voice_combo.setToolTip("Default narration voice")
        voice_layout.addWidget(self.default_voice_combo)

        self.coqui_check = QCheckBox("Enable fallback TTS")
        self.coqui_check.setChecked(self.app_settings.data.get("use_coqui_fallback", DEFAULTS["use_coqui_fallback"]))
        self.coqui_check.toggled.connect(lambda v: self.update_setting("use_coqui_fallback", v))
        self.coqui_check.setToolTip("Use Coqui if ElevenLabs fails")
        voice_layout.addWidget(self.coqui_check)

        # AI Settings ------------------------------------------------------
        ai_layout = section("\U0001F9E0 AI Settings")

        self.prompt_edit = QPlainTextEdit()
        self.prompt_edit.setPlaceholderText("Default AI prompt")
        self.prompt_edit.setPlainText(self.app_settings.data.get("ai_prompt", DEFAULTS["ai_prompt"]))
        self.prompt_edit.textChanged.connect(lambda: self.update_setting("ai_prompt", self.prompt_edit.toPlainText()))
        ai_layout.addWidget(self.prompt_edit)

        len_row = QHBoxLayout()
        self.len_spin = QSpinBox()
        self.len_spin.setRange(50, 1000)
        self.len_spin.setValue(self.app_settings.data.get("max_story_len", DEFAULTS["max_story_len"]))
        self.len_spin.valueChanged.connect(lambda v: self.update_setting("max_story_len", v))
        self.len_spin.setToolTip("Maximum story length")
        len_row.addWidget(QLabel("Max length"))
        len_row.addWidget(self.len_spin)
        len_row.addStretch()
        ai_layout.addLayout(len_row)

        # Buttons ----------------------------------------------------------
        btn_row = QHBoxLayout()
        self.restore_btn = PyPushButton(
            text="Restore Defaults",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["dark_three"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.restore_btn.clicked.connect(self.restore_defaults)
        btn_row.addWidget(self.restore_btn)

        self.save_btn = PyPushButton(
            text="Save Settings",
            radius=8,
            color=theme["text_foreground"],
            bg_color=theme["context_color"],
            bg_color_hover=theme["context_hover"],
            bg_color_pressed=theme["context_pressed"],
        )
        self.save_btn.clicked.connect(self.save_settings)
        btn_row.addWidget(self.save_btn)
        btn_row.addStretch()
        vbox.addLayout(btn_row)

        layout.addWidget(scroll)

    # ------------------------------------------------------------------
    def build_help_page(self):
        layout = self.ui.load_pages.page_3_layout
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        scroll = QScrollArea()
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setWidgetResizable(True)
        cont = QWidget()
        scroll.setWidget(cont)
        vbox = QVBoxLayout(cont)
        vbox.setAlignment(Qt.AlignTop)

        def add_section(title: str, text: str):
            lbl = QLabel(title)
            lbl.setObjectName("section")
            lbl.setStyleSheet(SECTION_STYLE)
            lbl.setToolTip(title)
            vbox.addWidget(lbl)
            body = QLabel(text)
            body.setWordWrap(True)
            vbox.addWidget(body)
            vbox.addSpacing(10)

        add_section(
            "\U0001F4D6 Getting Started",
            "Write or drop your story on the Home page, adjust options, then press 'Create Content'.",
        )
        add_section(
            "\U0001F3A5 Output Overview",
            "Videos are saved in the output folder using the title plus a timestamp. They work on TikTok, Shorts and Reels.",
        )
        add_section(
            "\u2699\ufe0f Troubleshooting",
            "If voiceover fails or subtitles are missing, check your settings and logs.",
        )

        btn = PyPushButton(
            text="Copy Logs",
            radius=8,
            color=self.themes["app_color"]["text_foreground"],
            bg_color=self.themes["app_color"]["dark_three"],
            bg_color_hover=self.themes["app_color"]["context_hover"],
            bg_color_pressed=self.themes["app_color"]["context_pressed"],
        )
        btn.clicked.connect(self.copy_logs)
        vbox.addWidget(btn, alignment=Qt.AlignLeft)
        vbox.addSpacing(10)

        layout.addWidget(scroll)

    # ------------------------------------------------------------------
    def set_status(self, text: str, color: str | None = None):
        """Update bottom status bar and reset after 5 seconds."""
        self.status_text.setText(text)
        if color:
            self.status_dot.setStyleSheet(f"color: {color}")
        effect = QGraphicsOpacityEffect(self.status_text)
        self.status_text.setGraphicsEffect(effect)
        anim = QPropertyAnimation(effect, b"opacity", self)
        anim.setDuration(200)
        anim.setStartValue(0)
        anim.setEndValue(1)
        anim.start(QPropertyAnimation.DeleteWhenStopped)
        self.status_timer.start(5000)

    # ------------------------------------------------------------------
    def update_setting(self, key: str, value):
        self.app_settings.data[key] = value
        self.app_settings.save()

    # ------------------------------------------------------------------
    def menu_clicked(self):
        pass

    def menu_released(self):
        pass

    def upload_script(self):
        path, _ = QFileDialog.getOpenFileName(self, "Open Script", "", "Text Files (*.txt *.docx)")
        if path:
            if path.lower().endswith(".txt"):
                with open(path, "r", encoding="utf-8") as f:
                    self.script_edit.setPlainText(f.read())
            elif path.lower().endswith(".docx"):
                doc = Document(path)
                text = "\n".join(p.text for p in doc.paragraphs)
                self.script_edit.setPlainText(text)
            else:
                QMessageBox.warning(self, "Unsupported", "Only .txt and .docx files are supported")
                return
            self.on_script_loaded(Path(path).name)

    def choose_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Output Folder", self.output_edit.text())
        if folder:
            path = Path(folder)
            if not path.exists():
                QMessageBox.warning(self, "Folder Missing", "Selected directory does not exist")
                return
            self.output_edit.setText(str(path))
            self.update_setting("output_folder", str(path))

    def surprise_me(self):
        combos = [self.voice_combo, self.subtitle_combo, self.bg_combo]
        for combo in combos:
            if combo.count():
                combo.setCurrentIndex(random.randrange(combo.count()))

    def generate_ai_story(self):
        prompt = (
            "Generate a 150-200 word horror story tailored for social media "
            "(TikTok/Shorts)."
        )
        theme = self.themes["app_color"]
        self.ai_btn.setEnabled(False)
        self.set_status("Generating AI...", theme.get("yellow"))
        QApplication.setOverrideCursor(Qt.WaitCursor)
        try:
            text = generator.generate_story(prompt=prompt)
            if not text:
                raise RuntimeError("No response from model")
            self.script_edit.setPlainText(text)
        except Exception as e:
            QMessageBox.critical(self, "AI Error", f"Story generation failed: {e}")
            self.set_status(f"\u274C Error: {e}", theme.get("red"))
        else:
            self.set_status("\u2705 Story Generated", theme.get("green"))
        finally:
            QApplication.restoreOverrideCursor()
            self.ai_btn.setEnabled(True)
            self.update_create_btn()

    def copy_logs(self):
        log_path = Path("logs/pipeline.log").resolve()
        if not log_path.exists():
            QMessageBox.warning(self, "Logs", "Log file not found")
            return
        QGuiApplication.clipboard().setText(str(log_path))
        self.set_status("Log path copied", self.themes["app_color"].get("green"))

    def on_script_loaded(self, name: str):
        self.set_status(f"Script loaded: {name}", self.themes["app_color"].get("green"))
        self.auto_filename = True
        self.on_script_changed()

    def disable_auto_title(self):
        self.auto_filename = False

    def on_script_changed(self):
        self.update_create_btn()
        if self.auto_filename:
            for line in self.script_edit.toPlainText().splitlines():
                if line.strip():
                    name = sanitize_name(line)[:50]
                    self.title_edit.blockSignals(True)
                    self.title_edit.setText(name)
                    self.title_edit.blockSignals(False)
                    break

    def sync_voice_combo(self, voice: str):
        i = self.voice_combo.findText(voice)
        if i >= 0:
            self.voice_combo.setCurrentIndex(i)

    def preview_voice(self):
        print(f"Previewing {self.voice_combo.currentText()}")

    def reset_form(self):
        self.script_edit.clear()
        self.title_edit.clear()
        for combo in (self.voice_combo, self.subtitle_combo, self.bg_combo):
            combo.setCurrentIndex(0)
        self.watermark_toggle.setChecked(False)
        self.auto_filename = True
        self.update_create_btn()
        self.update_watermark_label(False)

    def update_watermark_label(self, checked: bool):
        state = "On" if checked else "Off"
        self.watermark_label.setText(f"Watermark: {state}")
        self.output_info.setText(f"1080p @30fps | Watermark {state}")

    def update_create_btn(self):
        text = self.script_edit.toPlainText().strip()
        self.create_btn.setEnabled(bool(text))

    def run_pipeline(self):
        theme = self.themes["app_color"]
        self.create_btn.setEnabled(False)
        self.ai_btn.setEnabled(False)
        self.set_status("Generating Video...", theme.get("blue"))
        QApplication.setOverrideCursor(Qt.WaitCursor)
        try:
            cfg = Config.load(Path("config/config.json"))
            cfg.subtitle_style = self.subtitle_combo.currentText()
            cfg.default_voice_id = self.voice_combo.currentText()
            cfg.watermark_enabled = self.watermark_toggle.isChecked()
            cfg.validate()

            vp = VideoPipeline(cfg)
            title = self.title_edit.text().strip() or "session"
            script = self.script_edit.toPlainText()
            output_dir = Path(self.output_edit.text())
            out_path = output_dir / f"{sanitize_name(title)}.mp4"
            ctx = vp.run(
                script,
                title,
                background=self.bg_combo.currentText(),
                output=out_path,
                force_coqui=self.coqui_toggle.isChecked(),
            )
            QMessageBox.information(
                self,
                "Pipeline",
                f"Video saved to {ctx.final_video_path}",
            )
            self.set_status("\u2705 Content Created", theme.get("green"))
        except Exception as e:
            QMessageBox.critical(self, "Pipeline Error", str(e))
            self.set_status(f"\u274C Error: {e}", theme.get("red"))
        finally:
            QApplication.restoreOverrideCursor()
            self.create_btn.setEnabled(True)
            self.ai_btn.setEnabled(True)
            self.update_create_btn()

    def save_settings(self):
        d = self.app_settings.data
        d["output_folder"] = self.output_edit.text()
        d["use_coqui_fallback"] = self.coqui_check.isChecked()
        d["output_resolution"] = self.resolution_combo.currentText()
        d["watermark"] = self.wm_check.isChecked()
        d["ai_prompt"] = self.prompt_edit.toPlainText()
        d["max_story_len"] = self.len_spin.value()
        d["last_voice"] = self.default_voice_combo.currentText()
        d["last_subtitle"] = self.subtitle_combo.currentText()
        d["last_background"] = self.bg_combo.currentText()
        d["last_title"] = self.title_edit.text()
        self.app_settings.save()

    def restore_state(self):
        data = self.app_settings.data
        self.output_edit.setText(data.get("output_folder", DEFAULTS["output_folder"]))
        self.coqui_check.setChecked(data.get("use_coqui_fallback", DEFAULTS["use_coqui_fallback"]))
        self.resolution_combo.setCurrentText(data.get("output_resolution", DEFAULTS["output_resolution"]))
        self.wm_check.setChecked(data.get("watermark", DEFAULTS["watermark"]))
        self.prompt_edit.setPlainText(data.get("ai_prompt", DEFAULTS["ai_prompt"]))
        self.len_spin.setValue(data.get("max_story_len", DEFAULTS["max_story_len"]))
        voice = data.get("last_voice", DEFAULTS["last_voice"]) 
        idx = self.default_voice_combo.findText(voice)
        if idx >= 0:
            self.default_voice_combo.setCurrentIndex(idx)
        # home page widgets
        for combo, key in [
            (self.voice_combo, "last_voice"),
            (self.subtitle_combo, "last_subtitle"),
            (self.bg_combo, "last_background"),
        ]:
            val = data.get(key)
            if val is not None:
                i = combo.findText(val)
                if i >= 0:
                    combo.setCurrentIndex(i)
        self.watermark_toggle.setChecked(data.get("watermark", DEFAULTS["watermark"]))
        last_title = data.get("last_title", "")
        self.title_edit.setText(last_title)
        self.auto_filename = not bool(last_title)

    def restore_defaults(self):
        for k, v in DEFAULTS.items():
            self.app_settings.data[k] = v
        self.app_settings.save()
        self.restore_state()
        self.set_status("Defaults restored", self.themes["app_color"].get("green"))

    def adjust_preview_size(self):
        if not hasattr(self, "preview_container"):
            return
        frame_width = self.ui.load_pages.preview_frame.width()
        width = int(max(200, min(self.width() * 0.4, frame_width)))
        height = int(width * 16 / 9)
        self.preview_container.setFixedSize(width, height)

    def closeEvent(self, event):
        self.save_settings()
        super().closeEvent(event)

    # Resize grips
    def resizeEvent(self, event):
        if self.settings["custom_title_bar"]:
            self.left_grip.setGeometry(5, 10, 10, self.height())
            self.right_grip.setGeometry(self.width() - 15, 10, 10, self.height())
            self.top_grip.setGeometry(5, 5, self.width() - 10, 10)
            self.bottom_grip.setGeometry(5, self.height() - 15, self.width() - 10, 10)
            self.top_right_grip.setGeometry(self.width() - 20, 5, 15, 15)
            self.bottom_left_grip.setGeometry(5, self.height() - 20, 15, 15)
            self.bottom_right_grip.setGeometry(self.width() - 20, self.height() - 20, 15, 15)
        self.adjust_preview_size()
        super().resizeEvent(event)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    sys.exit(app.exec())

